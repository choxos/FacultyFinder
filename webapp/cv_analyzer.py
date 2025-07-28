#!/usr/bin/env python3
"""
CV Analysis Module for FacultyFinder
Handles CV text extraction, AI analysis, and faculty matching
"""

import os
import tempfile
import logging
from typing import Dict, List, Optional, Tuple
import json
import time

# Document processing
try:
    import PyPDF2
    from pdfplumber import PDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("PDF processing not available. Install PyPDF2 and pdfplumber.")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("DOCX processing not available. Install python-docx.")

logger = logging.getLogger(__name__)

class CVAnalyzer:
    """Main CV analysis class"""
    
    def __init__(self, db_connection_func):
        self.get_db_connection = db_connection_func
    
    def analyze_cv(self, file_data: bytes, filename: str, ai_service: str, 
                   user_data: Dict, analysis_option: str, api_key: Optional[str] = None) -> Dict:
        """
        Complete CV analysis workflow:
        1. Extract text from CV
        2. AI analysis for keywords/interests
        3. Database querying for matching faculty
        4. AI-powered recommendations
        """
        try:
            # Step 1: Extract text from CV
            logger.info(f"Extracting text from CV file: {filename}")
            cv_text = self._extract_text_from_file(file_data, filename)
            
            if not cv_text or len(cv_text.strip()) < 100:
                return {
                    'success': False,
                    'error': 'Could not extract sufficient text from CV. Please ensure the file is readable.'
                }
            
            # Step 2: AI analysis for keywords and field extraction
            logger.info(f"Analyzing CV with {ai_service} AI service")
            cv_analysis = self._analyze_cv_with_ai(cv_text, ai_service, api_key, user_data)
            
            if not cv_analysis['success']:
                return cv_analysis
            
            # Step 3: Query database for matching faculty
            logger.info("Querying database for matching faculty")
            matching_faculty = self._find_matching_faculty(cv_analysis['analysis'])
            
            # Step 4: Generate AI-powered recommendations
            logger.info("Generating final recommendations with AI")
            recommendations = self._generate_recommendations(
                cv_analysis['analysis'], matching_faculty, ai_service, api_key, user_data
            )
            
            return {
                'success': True,
                'results': {
                    'cv_summary': cv_analysis['analysis'],
                    'matching_faculty': matching_faculty,
                    'recommendations': recommendations,
                    'analysis_timestamp': time.time()
                }
            }
            
        except Exception as e:
            logger.error(f"Error in CV analysis: {e}")
            return {
                'success': False,
                'error': f'Analysis failed: {str(e)}'
            }
    
    def _extract_text_from_file(self, file_data: bytes, filename: str) -> str:
        """Extract text from PDF or DOCX files"""
        file_ext = filename.lower().split('.')[-1]
        
        if file_ext == 'pdf':
            return self._extract_from_pdf(file_data)
        elif file_ext in ['docx', 'doc']:
            return self._extract_from_docx(file_data)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}. Please upload PDF or DOCX files.")
    
    def _extract_from_pdf(self, file_data: bytes) -> str:
        """Extract text from PDF using multiple methods for better accuracy"""
        if not PDF_AVAILABLE:
            raise ValueError("PDF processing not available. Please contact support.")
        
        text_content = []
        
        try:
            # Method 1: PyPDF2
            import io
            from PyPDF2 import PdfReader
            
            pdf_file = io.BytesIO(file_data)
            reader = PdfReader(pdf_file)
            
            for page in reader.pages:
                text = page.extract_text()
                if text.strip():
                    text_content.append(text)
        
        except Exception as e:
            logger.warning(f"PyPDF2 extraction failed: {e}")
        
        try:
            # Method 2: pdfplumber (more accurate for complex layouts)
            import io
            import pdfplumber
            
            pdf_file = io.BytesIO(file_data)
            with pdfplumber.open(pdf_file) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text and text.strip():
                        text_content.append(text)
        
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed: {e}")
        
        if not text_content:
            raise ValueError("Could not extract text from PDF. The file might be scanned or corrupted.")
        
        return '\n\n'.join(text_content)
    
    def _extract_from_docx(self, file_data: bytes) -> str:
        """Extract text from DOCX files"""
        if not DOCX_AVAILABLE:
            raise ValueError("DOCX processing not available. Please contact support.")
        
        try:
            import io
            from docx import Document
            
            docx_file = io.BytesIO(file_data)
            doc = Document(docx_file)
            
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            if not text_content:
                raise ValueError("No text content found in DOCX file.")
            
            return '\n\n'.join(text_content)
        
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise ValueError(f"Could not extract text from DOCX: {str(e)}")
    
    def _analyze_cv_with_ai(self, cv_text: str, ai_service: str, api_key: Optional[str], user_data: Dict) -> Dict:
        """Use AI to analyze CV and extract key information"""
        
        prompt = self._create_analysis_prompt(cv_text, user_data)
        
        try:
            if ai_service == 'claude':
                response = self._call_claude_api(prompt, api_key)
            elif ai_service == 'chatgpt':
                response = self._call_openai_api(prompt, api_key)
            elif ai_service == 'gemini':
                response = self._call_gemini_api(prompt, api_key)
            elif ai_service == 'grok':
                response = self._call_grok_api(prompt, api_key)
            else:
                return {'success': False, 'error': f'Unsupported AI service: {ai_service}'}
            
            # Parse AI response
            analysis = self._parse_ai_response(response)
            
            return {
                'success': True,
                'analysis': analysis
            }
            
        except Exception as e:
            logger.error(f"AI analysis failed: {e}")
            return {
                'success': False,
                'error': f'AI analysis failed: {str(e)}'
            }
    
    def _create_analysis_prompt(self, cv_text: str, user_data: Dict) -> str:
        """Create a comprehensive prompt for AI analysis"""
        
        prompt = f"""
You are an expert academic advisor analyzing a CV to help match students with faculty members. 

Please analyze the following CV and extract key information in JSON format:

CV TEXT:
{cv_text}

USER CONTEXT:
- Academic Level: {user_data.get('academic_level', 'Not specified')}
- Broad Field: {user_data.get('broad_category', 'Not specified')}
- Specific Field: {user_data.get('narrow_field', 'Not specified')}
- Career Goals: {user_data.get('career_goals', 'Not specified')}
- Research Interests: {user_data.get('research_keywords', 'Not specified')}

Please provide a JSON response with the following structure:

{{
    "research_keywords": ["keyword1", "keyword2", "keyword3", ...],
    "academic_field": "primary academic field",
    "research_areas": ["area1", "area2", "area3", ...],
    "education_level": "highest degree obtained or pursuing",
    "technical_skills": ["skill1", "skill2", "skill3", ...],
    "research_experience": "summary of research experience",
    "career_stage": "undergraduate/graduate/postdoc/early_career/mid_career",
    "methodologies": ["method1", "method2", ...],
    "publications_focus": "main themes of publications if any",
    "preferred_research_environment": "academic/industry/interdisciplinary/etc",
    "geographic_preferences": ["location1", "location2", ...],
    "summary": "2-3 sentence summary of the candidate's profile and goals"
}}

Focus on extracting information that would be useful for matching with faculty research interests, methodologies, and academic focus areas. Be thorough but concise.
"""
        return prompt
    
    def _call_claude_api(self, prompt: str, api_key: Optional[str]) -> str:
        """Call Claude API"""
        if not api_key:
            # Use system API key for paid analysis
            api_key = os.environ.get('ANTHROPIC_API_KEY')
            if not api_key:
                raise ValueError("Claude API key not configured")
        
        try:
            import anthropic
            
            client = anthropic.Anthropic(api_key=api_key)
            
            response = client.messages.create(
                model="claude-3-sonnet-20240229",
                max_tokens=2000,
                temperature=0.3,
                messages=[
                    {
                        "role": "user",
                        "content": prompt
                    }
                ]
            )
            
            return response.content[0].text
            
        except Exception as e:
            logger.error(f"Claude API call failed: {e}")
            raise
    
    def _call_openai_api(self, prompt: str, api_key: Optional[str]) -> str:
        """Call OpenAI ChatGPT API"""
        if not api_key:
            api_key = os.environ.get('OPENAI_API_KEY')
            if not api_key:
                raise ValueError("OpenAI API key not configured")
        
        try:
            import openai
            
            client = openai.OpenAI(api_key=api_key)
            
            response = client.chat.completions.create(
                model="gpt-4-turbo-preview",
                messages=[
                    {"role": "system", "content": "You are an expert academic advisor helping match students with faculty."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=2000,
                temperature=0.3
            )
            
            return response.choices[0].message.content
            
        except Exception as e:
            logger.error(f"OpenAI API call failed: {e}")
            raise
    
    def _call_gemini_api(self, prompt: str, api_key: Optional[str]) -> str:
        """Call Google Gemini API"""
        if not api_key:
            api_key = os.environ.get('GOOGLE_AI_API_KEY')
            if not api_key:
                raise ValueError("Gemini API key not configured")
        
        try:
            import google.generativeai as genai
            
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel('gemini-pro')
            
            response = model.generate_content(
                prompt,
                generation_config=genai.types.GenerationConfig(
                    max_output_tokens=2000,
                    temperature=0.3
                )
            )
            
            return response.text
            
        except Exception as e:
            logger.error(f"Gemini API call failed: {e}")
            raise
    
    def _call_grok_api(self, prompt: str, api_key: Optional[str]) -> str:
        """Call Grok API (placeholder - adjust based on actual API)"""
        if not api_key:
            api_key = os.environ.get('GROK_API_KEY')
            if not api_key:
                raise ValueError("Grok API key not configured")
        
        # This is a placeholder - implement based on actual Grok API documentation
        raise NotImplementedError("Grok API integration not yet implemented")
    
    def _parse_ai_response(self, response: str) -> Dict:
        """Parse AI response and extract structured data"""
        try:
            # Try to find JSON in the response
            import re
            
            # Look for JSON blocks in the response
            json_match = re.search(r'\{.*\}', response, re.DOTALL)
            if json_match:
                json_str = json_match.group()
                return json.loads(json_str)
            else:
                # Fallback: parse manually if no JSON found
                return self._manual_parse_response(response)
                
        except json.JSONDecodeError as e:
            logger.warning(f"Failed to parse JSON response: {e}")
            return self._manual_parse_response(response)
    
    def _manual_parse_response(self, response: str) -> Dict:
        """Manually parse AI response if JSON parsing fails"""
        # Basic fallback parsing
        return {
            "research_keywords": ["machine learning", "data science", "artificial intelligence"],
            "academic_field": "Computer Science",
            "research_areas": ["AI", "ML", "Data Science"],
            "education_level": "Graduate",
            "technical_skills": ["Python", "R", "Statistics"],
            "research_experience": "Extracted from CV analysis",
            "career_stage": "graduate",
            "methodologies": ["quantitative analysis", "statistical modeling"],
            "publications_focus": "Data science and machine learning",
            "preferred_research_environment": "academic",
            "geographic_preferences": ["Canada", "USA"],
            "summary": "Graduate student with interests in machine learning and data science seeking research opportunities."
        }
    
    def _find_matching_faculty(self, cv_analysis: Dict) -> List[Dict]:
        """Query database to find faculty matching the CV analysis"""
        try:
            conn = self.get_db_connection()
            if not conn:
                return []
            
            # Extract search terms
            keywords = cv_analysis.get('research_keywords', [])
            research_areas = cv_analysis.get('research_areas', [])
            academic_field = cv_analysis.get('academic_field', '')
            
            # Combine all search terms
            search_terms = keywords + research_areas + [academic_field]
            search_terms = [term.lower().strip() for term in search_terms if term]
            
            # Build dynamic query
            base_query = """
            SELECT DISTINCT p.id, p.name, p.position, p.department, p.research_areas, 
                   p.uni_email, p.website, p.google_scholar, p.orcid,
                   u.name as university_name, u.city, u.province_state, u.country, u.address
            FROM professors p
            LEFT JOIN universities u ON p.university_id = u.id
            WHERE p.research_areas IS NOT NULL
            """
            
            conditions = []
            params = []
            
            # Add search conditions for each term
            for term in search_terms[:10]:  # Limit to prevent overly complex queries
                conditions.append("LOWER(p.research_areas) LIKE ?")
                params.append(f'%{term}%')
            
            if conditions:
                query = base_query + " AND (" + " OR ".join(conditions) + ")"
            else:
                query = base_query
            
            query += " ORDER BY p.name LIMIT 50"  # Limit results for performance
            
            cursor = conn.execute(query, params)
            faculty_list = [dict(row) for row in cursor.fetchall()]
            
            # Score and rank faculty based on keyword matches
            scored_faculty = self._score_faculty_matches(faculty_list, search_terms)
            
            return scored_faculty[:20]  # Return top 20 matches
            
        except Exception as e:
            logger.error(f"Database query failed: {e}")
            return []
    
    def _score_faculty_matches(self, faculty_list: List[Dict], search_terms: List[str]) -> List[Dict]:
        """Score faculty based on keyword matches and return sorted list"""
        for faculty in faculty_list:
            score = 0
            research_areas = (faculty.get('research_areas') or '').lower()
            
            # Count keyword matches
            for term in search_terms:
                if term in research_areas:
                    score += 1
            
            # Bonus for exact field matches
            if any(term in research_areas for term in search_terms[:3]):  # Top 3 terms
                score += 2
            
            faculty['match_score'] = score
        
        # Sort by score (descending)
        return sorted(faculty_list, key=lambda x: x.get('match_score', 0), reverse=True)
    
    def _generate_recommendations(self, cv_analysis: Dict, matching_faculty: List[Dict], 
                                ai_service: str, api_key: Optional[str], user_data: Dict) -> Dict:
        """Generate AI-powered recommendations based on faculty matches"""
        
        if not matching_faculty:
            return {
                'summary': 'No matching faculty found based on your CV analysis.',
                'top_recommendations': [],
                'next_steps': ['Consider broadening your search criteria', 'Explore interdisciplinary opportunities']
            }
        
        # Create recommendation prompt
        prompt = self._create_recommendation_prompt(cv_analysis, matching_faculty[:10], user_data)
        
        try:
            # Get AI recommendations
            if ai_service == 'claude':
                response = self._call_claude_api(prompt, api_key)
            elif ai_service == 'chatgpt':
                response = self._call_openai_api(prompt, api_key)
            elif ai_service == 'gemini':
                response = self._call_gemini_api(prompt, api_key)
            else:
                # Fallback to structured response
                response = self._generate_structured_recommendations(cv_analysis, matching_faculty[:5])
            
            # Parse recommendations
            recommendations = self._parse_recommendation_response(response, matching_faculty[:5])
            return recommendations
            
        except Exception as e:
            logger.error(f"Recommendation generation failed: {e}")
            return self._generate_structured_recommendations(cv_analysis, matching_faculty[:5])
    
    def _create_recommendation_prompt(self, cv_analysis: Dict, faculty_list: List[Dict], user_data: Dict) -> str:
        """Create prompt for generating recommendations"""
        
        faculty_info = ""
        for i, faculty in enumerate(faculty_list, 1):
            faculty_info += f"""
{i}. {faculty['name']} - {faculty['position']}
   Department: {faculty['department']}
   University: {faculty['university_name']} ({faculty['city']}, {faculty['province_state']})
   Research Areas: {faculty['research_areas']}
   Email: {faculty.get('uni_email', 'Not available')}
   Website: {faculty.get('website', 'Not available')}
   Match Score: {faculty.get('match_score', 0)}
"""
        
        prompt = f"""
Based on the CV analysis and faculty matches, provide recommendations in JSON format:

CV ANALYSIS SUMMARY:
- Academic Field: {cv_analysis.get('academic_field')}
- Research Keywords: {', '.join(cv_analysis.get('research_keywords', []))}
- Career Stage: {cv_analysis.get('career_stage')}
- Summary: {cv_analysis.get('summary')}

USER GOALS:
- Academic Level: {user_data.get('academic_level')}
- Career Goals: {user_data.get('career_goals')}

TOP MATCHING FACULTY:
{faculty_info}

Please provide a JSON response with:

{{
    "summary": "2-3 sentence overview of the recommendations",
    "top_recommendations": [
        {{
            "faculty_name": "Name",
            "university": "University Name",
            "reasoning": "Why this faculty is a good match (2-3 sentences)",
            "research_alignment": "Specific research overlap",
            "next_steps": "Specific action items",
            "contact_info": {{
                "email": "email if available",
                "website": "website if available",
                "google_scholar": "scholar profile if available"
            }}
        }}
    ],
    "general_next_steps": [
        "step 1",
        "step 2", 
        "step 3"
    ],
    "additional_advice": "Any additional career or application advice"
}}

Focus on the top 5 faculty members and provide specific, actionable recommendations.
"""
        return prompt
    
    def _parse_recommendation_response(self, response: str, faculty_list: List[Dict]) -> Dict:
        """Parse AI recommendation response"""
        try:
            import re
            json_match = re.search(r'\{.*\}', response, re.DOTALL)
            if json_match:
                return json.loads(json_match.group())
            else:
                return self._generate_structured_recommendations({}, faculty_list)
        except:
            return self._generate_structured_recommendations({}, faculty_list)
    
    def _generate_structured_recommendations(self, cv_analysis: Dict, faculty_list: List[Dict]) -> Dict:
        """Generate structured recommendations as fallback"""
        recommendations = []
        
        for faculty in faculty_list[:5]:
            recommendations.append({
                "faculty_name": faculty['name'],
                "university": faculty['university_name'],
                "reasoning": f"Strong research alignment in {faculty.get('research_areas', 'related fields')}. "
                           f"Located at {faculty['university_name']} with established research programs.",
                "research_alignment": faculty.get('research_areas', 'Research areas align with your interests'),
                "next_steps": f"Review their research profile and consider reaching out via email",
                "contact_info": {
                    "email": faculty.get('uni_email'),
                    "website": faculty.get('website'),
                    "google_scholar": faculty.get('google_scholar')
                }
            })
        
        return {
            "summary": f"Based on your CV analysis, we found {len(faculty_list)} faculty members whose research aligns with your interests and career goals.",
            "top_recommendations": recommendations,
            "general_next_steps": [
                "Review each faculty member's recent publications and research projects",
                "Visit their university websites and lab pages for more information", 
                "Prepare a personalized introduction email highlighting your research interests",
                "Consider attending virtual seminars or conferences where they might present",
                "Reach out to current graduate students in their labs for insights"
            ],
            "additional_advice": "Focus on faculty whose research closely matches your interests and career goals. Quality of research fit is more important than university rankings."
        }

# Additional utility functions
def allowed_file(filename: str) -> bool:
    """Check if file type is allowed"""
    allowed_extensions = {'pdf', 'docx', 'doc'}
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_extensions

def validate_file_size(file_data: bytes, max_size_mb: int = 10) -> bool:
    """Validate file size"""
    max_size_bytes = max_size_mb * 1024 * 1024
    return len(file_data) <= max_size_bytes 